"""Build the Matrix PoC status doc as a landscape-orientation .docx with narrow margins.
All body text avoids the ASCII hyphen character.
"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).resolve().parents[1] / "samples" / "matrix_status_update.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()

# Landscape + narrow margins on all sections.
for section in doc.sections:
    w, h = section.page_width, section.page_height
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = h, w
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

# Body style.
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)


def h2(text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def p(text: str, *, italic: bool = False) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.italic = italic
    para.paragraph_format.space_after = Pt(4)


def bullet(text: str, *, bold_lead: str | None = None) -> None:
    para = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r1 = para.add_run(bold_lead)
        r1.bold = True
        para.add_run(" " + text)
    else:
        para.add_run(text)
    para.paragraph_format.space_after = Pt(2)


# === Title block ===
h1("Matrix PoC \u2014 Status Update")
p("From: Anup Roy.  Context: response to manager note about doc extraction grid, "
  "embeddings gap, and Q&A framework.",
  italic=True)

# Replace the em dash (\u2014) in the title with a non hyphen separator — we allowed
# em dashes in the title only because the spec forbids the ASCII hyphen "-". If strict
# no dash at all, adjust here.
# Actually: the user wrote "without '-'". An em dash is a different glyph. Keeping title
# clean with a middle dot separator instead.
# (Swap above title programmatically.)
doc.paragraphs[0].clear()
run = doc.paragraphs[0].add_run("Matrix PoC \u00b7 Status Update")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

# === What we built ===
h2("1. What we built")
p("A working Hebbia Matrix style spreadsheet over PDFs. Rows are documents, columns are "
  "prompts, cells stream answers with inline citations that link back to the exact PDF "
  "page. The UI is a premium dark theme grid with a command palette, a split pane focus "
  "view that embeds a PDF viewer with bounding box highlights, and a synthesis dock for "
  "cross row aggregation.")

# === What powers it ===
h2("2. What powers it under the hood")

p("Parsing. Every PDF page is rendered to an image and sent to gpt 4.1 vision. The model "
  "returns clean markdown. This avoids the usual fragility of text extraction libraries "
  "when documents have tables, footnotes, or complex layouts.")

p("Embeddings and retrieval. Every chunk is embedded into LanceDB at ingest. We try Azure "
  "text embedding 3 large first, then automatically fall back to a local bge large English "
  "sentence transformer if that deployment is not exposed. Three retriever modes ship in "
  "the box, all behind the same interface so they are swappable at runtime:")

bullet("Vector top k. The baseline retriever.", bold_lead="Naive.")
bullet("Iterative Source Decomposition. Two LLM passes per query. Pass one decomposes the "
       "user intent into sub queries and identifies target sections from the document "
       "section index. Pass two does batched attention reranking, scoring every candidate "
       "chunk on a zero to one relevance scale in a single model call. This is the Hebbia "
       "signature move.", bold_lead="ISD.")
bullet("The tier three idea. During ingest we build a structured per document wiki: "
       "section summaries, named entities, notable claims with evidence chunk ids, "
       "quantitative metrics keyed to chunk ids, and a doc level rollup with an overview "
       "and a key metrics table. Retrieval queries this structured knowledge first and "
       "falls back to vector search only when the wiki does not cover the question. "
       "Closer to how a human remembers a document than plain vector search.",
       bold_lead="Wiki.")

p("ISD agent per cell. Each cell runs a decompose, draft, verify, revise loop. The agent "
  "retrieves evidence per sub question, drafts an answer with cited chunks, verifies each "
  "claim by re retrieving and asking the model whether the evidence supports the claim, "
  "and revises once if anything is contradicted or missing. Every step is logged to a "
  "gzipped trace on disk.")

p("Synthesis. A bottom dock accepts a free form question across the filled grid and "
  "produces a narrative that cites back to specific row and column cells.")

# === Addresses ===
h2("3. How this addresses the two items you flagged")

p("On the embeddings concern. We are already past raw character truncation. Vector "
  "retrieval with bge large is live, and the ISD retriever layers prompt decomposition, "
  "section targeting, and batched attention reranking on top of it. LanceDB stores "
  "embeddings per document in an isolated table keyed by the PDF hash.")

p("On the ask questions about your data framework. That is exactly the per cell ISD agent. "
  "It is live, streaming state through server sent events, with citations that open the "
  "source PDF in a pane on the right side of the screen. A FinanceBench benchmark harness "
  "runs all three retriever modes against the same questions and emits a markdown "
  "comparison table with correctness, citation page recall, citation precision, latency, "
  "and token cost per question.")

# === Current state ===
h2("4. Current state")
bullet("Twenty three backend tests passing. Frontend type check clean.")
bullet("End to end demo working on a sample ten K with real answers and citations that "
       "highlight the bounding box on the PDF page.")
bullet("FinanceBench harness plugs in with one command and produces a per mode report.")
bullet("SQLite database with WAL mode, LanceDB vector index, gzipped per cell traces, "
       "content addressable PDF storage.")

# === Next ===
h2("5. Next")
bullet("Real scale. Run the benchmark on a full FinanceBench slice and publish the three "
       "way comparison.")
bullet("Tighter citations. Extend bounding box resolution from page level to paragraph "
       "level so the overlay pinpoints the sentence.")
bullet("Wiki schema versioning. Lock the current schema and add a rebuild on bump path so "
       "production wikis can evolve without breaking historical benchmarks.")
bullet("Cost guardrail. Add a token budget preview before large column fan outs.")

# === Footer ===
p(" ")
p("Repo: /Users/anup.roy/Downloads/Hebbia POC  |  Spec: docs/superpowers/specs/  |  "
  "Plan: docs/superpowers/plans/  |  Architecture diagrams: docs/architecture.md",
  italic=True)

doc.save(OUT)
print(f"wrote {OUT}  ({OUT.stat().st_size} bytes)")
